import random
import docx
from googletrans import Translator

file_path = "documents/rewiew.docx"
name = "documents/" + str(random.random()) + ".docx"
translator = Translator()

doc = docx.Document()
file = docx.Document(file_path)
paragraphs = file.paragraphs
for paragraph in paragraphs:
    print(paragraph.text)
    if paragraph.text != "":
        translated_text = translator.translate(paragraph.text, src='russian')
        print(translated_text.text)
        doc.add_paragraph(translated_text.text, paragraph.style)
name = 'documents/' + str(random.random()) + ".docx"
doc.save(name)


