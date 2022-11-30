import docx
from googletrans import Translator

translator = Translator()


# Class Translator here
class TranslatorCl:
    from_lang = None  # first language
    to_lang = None  # second language

# The Constructor here
    def __init__(self):
        self.from_lang = None
        self.to_lang = None

# Setting the first(home) language
    def set_from_lan(self, fl):
        self.from_lang = fl

# Setting the second language(to translate in)
    def set_to_lan(self, tl):
        self.to_lang = tl

# Translation of a text
    def translate(self, text):
        result = translator.translate(text, src=self.from_lang,
                                      dest=self.to_lang)
        return result.text

# The method translates .docx files
    def translate_docx(self, file_path):
        doc = docx.Document()  # create new docx file
        file = docx.Document(file_path)  # open the file to translate
        paragraphs = file.paragraphs  # get all paragraphs
        for paragraph in paragraphs:
            if paragraph.text != '':  # check if paragraph is not empty
                translated_text = translator.translate(paragraph.text, src=self.from_lang,
                                                       dest=self.to_lang)  # translate paragraph
                doc.add_paragraph(translated_text.text, paragraph.style.name)  # add translated paragraph to new file
            else:
                doc.add_paragraph(paragraph.text, paragraph.style.name)  # add empty paragraph to new file
        name = 'documents/Translated_text.docx'  # name of the new file
        doc.save(name)  # save new file
        return name

# Translates .txt files
# Just read translating file, translate its text, write translated text in new file
# Save translated file and return its path
    def translate_txt(self, file_path):
        name = 'documents/Translated_text.txt'  # name of the new file
        with open(name, 'w') as translated_file, open(file_path, 'r') as file:  # open files
            text = file.read()  # read text from file
            translated_text = translator.translate(text, src=self.from_lang,
                                                   dest=self.to_lang)  # translate text
            translated_file.write(translated_text.text)  # write translated text in new file
            return name

    def translate_for_chat(self, text):
        text = text.lower()  # make text lower case
        try:  # try to translate text
            result = translator.translate(text, src=self.from_lang,
                                          dest=self.to_lang)  # translate text
            # check if text is not in the first language
            if result.text.lower() == text or translator.detect(text).lang == self.to_lang:
                result = translator.translate(text, src=self.to_lang,
                                              dest=self.from_lang)  # translate text
            return result.text.lower()  # return translated text
        except Exception:  # if translation failed
            print("Chat_Translator Error - Try again after 5 secs")
            return "I can't translate this text:("
            # return error message

# The method to clear properties of the class
    def forget(self):
        self.from_lang = None
        self.to_lang = None
